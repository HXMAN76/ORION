"""DOCX document processor"""
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.table import Table
import io
import tempfile

from .base import BaseProcessor, Chunk
from ..chunking import Chunker
from ..ocr import DeepSeekOCR


class DOCXProcessor(BaseProcessor):
    """Process DOCX files extracting text, tables, and embedded images"""
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]
    
    @property
    def doc_type(self) -> str:
        return "docx"
    
    def __init__(self, chunker: Optional[Chunker] = None, use_ocr_for_images: bool = True):
        """
        Initialize DOCX processor.
        
        Args:
            chunker: Text chunker instance (uses default if not provided)
            use_ocr_for_images: Whether to OCR embedded images (enabled by default)
        """
        self.chunker = chunker or Chunker()
        self.use_ocr_for_images = use_ocr_for_images
        self._ocr = None
    
    def process(self, file_path: Path, document_id: Optional[str] = None) -> List[Chunk]:
        """
        Process a DOCX file and extract chunks.
        
        Args:
            file_path: Path to the DOCX file
            document_id: Optional document identifier
            
        Returns:
            List of Chunk objects with text content and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        document_id = document_id or self._generate_document_id(file_path)
        chunks: List[Chunk] = []
        
        doc = Document(file_path)
        
        # Extract paragraphs
        full_text = []
        current_section = ""
        
        for para in doc.paragraphs:
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                if full_text:
                    # Chunk the accumulated text
                    self._add_text_chunks(
                        chunks, 
                        "\n".join(full_text), 
                        document_id, 
                        file_path,
                        current_section
                    )
                    full_text = []
                current_section = para.text
            else:
                if para.text.strip():
                    full_text.append(para.text)
        
        # Don't forget the last section
        if full_text:
            self._add_text_chunks(
                chunks, 
                "\n".join(full_text), 
                document_id, 
                file_path,
                current_section
            )
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_md = self._table_to_markdown(table)
            if table_md:
                chunks.append(Chunk(
                    content=table_md,
                    document_id=document_id,
                    doc_type=self.doc_type,
                    source_file=str(file_path),
                    metadata={
                        "content_type": "table",
                        "table_index": table_idx
                    }
                ))
        
        # Extract embedded images and run OCR
        if self.use_ocr_for_images:
            image_chunks = self._extract_embedded_images(file_path, doc, document_id)
            chunks.extend(image_chunks)
        
        return chunks
    
    def _get_ocr(self):
        """Lazy load OCR engine"""
        if self._ocr is None and self.use_ocr_for_images:
            try:
                self._ocr = DeepSeekOCR()
                if not self._ocr.is_available():
                    print("Warning: DeepSeek model not found in Ollama, image OCR disabled")
                    print("To enable image OCR, run: ollama pull deepseek-vl")
                    self._ocr = None
            except Exception as e:
                print(f"Warning: DeepSeek OCR failed to initialize: {e}")
                self._ocr = None
        return self._ocr
    
    def _extract_embedded_images(
        self, 
        file_path: Path, 
        doc: Document, 
        document_id: str
    ) -> List[Chunk]:
        """Extract and OCR embedded images from DOCX"""
        chunks: List[Chunk] = []
        ocr = self._get_ocr()
        
        if ocr is None:
            return chunks
        
        try:
            # Extract images from document relationships
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        # Get image binary data
                        image_data = rel.target_part.blob
                        
                        # Save to temporary file
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                            tmp_path = Path(tmp.name)
                            tmp.write(image_data)
                        
                        try:
                            # Run OCR on the image
                            content = ocr.extract_full_content(tmp_path)
                            
                            # Add text chunk if available
                            if content.get("text", "").strip():
                                chunks.append(Chunk(
                                    content=f"[Embedded Image Text]\n{content['text']}",
                                    document_id=document_id,
                                    doc_type=self.doc_type,
                                    source_file=str(file_path),
                                    metadata={
                                        "content_type": "embedded_image_text",
                                        "extraction_method": "deepseek",
                                        "image_rel_id": rel_id
                                    }
                                ))
                            
                            # Add table chunks if available
                            for table_idx, table_md in enumerate(content.get("tables", [])):
                                chunks.append(Chunk(
                                    content=table_md,
                                    document_id=document_id,
                                    doc_type=self.doc_type,
                                    source_file=str(file_path),
                                    metadata={
                                        "content_type": "embedded_image_table",
                                        "table_index": table_idx,
                                        "extraction_method": "deepseek",
                                        "image_rel_id": rel_id
                                    }
                                ))
                        finally:
                            # Clean up temporary file
                            if tmp_path.exists():
                                tmp_path.unlink()
                    
                    except Exception as e:
                        print(f"Failed to process embedded image {rel_id}: {e}")
                        continue
        
        except Exception as e:
            print(f"Failed to extract embedded images: {e}")
        
        return chunks
    
    def _add_text_chunks(
        self, 
        chunks: List[Chunk], 
        text: str, 
        document_id: str,
        file_path: Path,
        section: str = ""
    ):
        """Add chunked text to the chunks list"""
        text_chunks = self.chunker.chunk(text)
        
        for idx, chunk_text in enumerate(text_chunks):
            content = chunk_text
            if section:
                content = f"[Section: {section}]\n\n{chunk_text}"
            
            chunks.append(Chunk(
                content=content,
                document_id=document_id,
                doc_type=self.doc_type,
                source_file=str(file_path),
                metadata={
                    "section": section,
                    "chunk_index": idx
                }
            ))
    
    def _table_to_markdown(self, table: Table) -> str:
        """Convert a DOCX table to markdown format"""
        rows = []
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) >= 1:
            # Add header separator
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
        
        return "\n".join(rows) if rows else ""
